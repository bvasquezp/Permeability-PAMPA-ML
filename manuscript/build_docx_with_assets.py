"""Create a Word manuscript version with Python-generated tables and figures."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


MANUSCRIPT = Path(__file__).resolve().parent
BASE_DOCX = MANUSCRIPT / "pampa_qsar_manuscript.docx"
OUTPUT_DOCX = MANUSCRIPT / "pampa_qsar_manuscript_with_python_assets.docx"
TABLES = MANUSCRIPT / "tables"
FIGURES = MANUSCRIPT / "figures"


def add_dataframe_table(document: Document, title: str, csv_name: str) -> None:
    document.add_heading(title, level=2)
    df = pd.read_csv(TABLES / csv_name)
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, column in enumerate(df.columns):
        cell = table.rows[0].cells[i]
        cell.text = str(column)
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row.tolist()):
            cells[i].text = str(value)
    document.add_paragraph("")


def add_figure(document: Document, title: str, filename: str, width: float = 5.8) -> None:
    document.add_heading(title, level=2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(FIGURES / filename), width=Inches(width))


def add_manuscript_additions(document: Document) -> None:
    document.add_heading("Expanded methods and interpretation added to the LaTeX manuscript", level=2)
    document.add_paragraph(
        "The updated manuscript now describes the alvaDesc v.2.0.10 calculation from SMILES, "
        "limited to 0D-2D descriptors plus MACCS/ECFP fingerprints. Constant, near-constant "
        "and highly correlated descriptors were removed in alvaDesc, leaving 783 variables."
    )
    document.add_paragraph(
        "The EDA stage is described as an explicit dimensionality-reduction step: integrity "
        "checks, class-balance inspection, variance filtering, Mann-Whitney and correlation "
        "analyses, and Random-Forest RFE reduced the descriptor space to 50 variables before "
        "WEKA consensus voting selected the final 11 descriptors."
    )
    document.add_paragraph(
        "The results section also includes SHAP interpretation, named prospective pre-screening "
        "controls, and a DrugBank funnel from more than 10000 source entries to 863 high-probability "
        "candidates and 770 final candidates that are both inside the applicability domain and "
        "Lipinski viable."
    )


def main() -> None:
    document = Document(BASE_DOCX)
    document.add_page_break()
    document.add_heading("Python-generated reproducibility tables and figures", level=1)
    document.add_paragraph(
        "The following tables and figures were generated directly from the repository "
        "datasets, model and result files using manuscript/generate_python_assets.py."
    )
    add_manuscript_additions(document)

    add_dataframe_table(document, "Generated Table S1. Dataset composition", "dataset_composition.csv")
    add_dataframe_table(document, "Generated Table S2. Cross-validation metrics", "cross_validation_metrics.csv")
    add_dataframe_table(document, "Generated Table S3. Final model metrics", "final_metrics.csv")
    add_dataframe_table(document, "Generated Table S4. Named pre-screening results", "pre_screening_named.csv")
    add_dataframe_table(document, "Generated Table S5. DrugBank filtering flow", "drugbank_filtering_flow.csv")
    add_dataframe_table(document, "Generated Table S6. DrugBank screening summary", "drugbank_screening_summary.csv")
    add_dataframe_table(document, "Generated Table S7. Top DrugBank candidates with names", "top_drugbank_candidates.csv")
    add_dataframe_table(document, "Generated Table S8. Thesis-reported DrugBank top 10 audit", "drugbank_thesis_reported_top10.csv")

    add_figure(document, "Generated Figure S1. Dataset class distribution", "python_dataset_class_distribution.png")
    add_figure(document, "Generated Figure S2. Final metrics heatmap", "python_final_metrics_heatmap.png")
    add_figure(document, "Generated Figure S3. Confusion matrices", "python_confusion_matrices.png", width=6.3)
    add_figure(document, "Generated Figure S4. SHAP global importance", "SHAP_Bar_PAMPA.png")
    add_figure(document, "Generated Figure S5. SHAP beeswarm interpretation", "SHAP_Beeswarm_PAMPA.png")
    add_figure(document, "Generated Figure S6. Pre-screening structures and controls", "python_pre_screening_structures.png", width=6.4)
    add_figure(document, "Generated Figure S7. Pre-screening probabilities", "python_pre_screening_probabilities.png")
    add_figure(document, "Generated Figure S8. DrugBank filtering funnel", "python_drugbank_filtering_flow.png")
    add_figure(document, "Generated Figure S9. DrugBank probability distribution", "python_drugbank_probability_distribution.png")
    add_figure(document, "Generated Figure S10. DrugBank leverage profile", "python_drugbank_leverage_probability.png")

    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX.resolve())


if __name__ == "__main__":
    main()
